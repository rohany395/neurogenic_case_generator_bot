import streamlit as st
import streamlit.components.v1 as components
import openai
from anthropic import Anthropic
from datetime import datetime
from typing import Optional
import json
import html
import os
import glob
import random

# Fix for SQLite version issue with ChromaDB
try:
    __import__('pysqlite3')
    import sys
    sys.modules['sqlite3'] = sys.modules.pop('pysqlite3')
except ImportError:
    pass  

# RAG imports
from langchain_community.document_loaders import TextLoader
from langchain_text_splitters import CharacterTextSplitter
from langchain_openai import OpenAIEmbeddings
from langchain_community.vectorstores import Chroma
from docx import Document as DocxDocument

LANGUAGE_PROFILE_HEADING_TEXT = "Language Profile / Communication Observations"
LANGUAGE_PROFILE_HEADING = f"### {LANGUAGE_PROFILE_HEADING_TEXT}"

# RAG Configuration
REF_DOCUMENTS_DIR = "ref_documents"
CHROMA_DB_DIR = "data/chroma_db"


def load_docx_content(file_path: str) -> str:
    """Extract text content from a .docx file."""
    try:
        doc = DocxDocument(file_path)
        return "\n".join(para.text for para in doc.paragraphs if para.text.strip())
    except Exception:
        return ""


def get_best_matching_exemplar(user_input: str, openai_key: Optional[str]) -> tuple[str, str]:
    """
    Load a reference case from ref_documents.
    Uses semantic search if vector DB exists, otherwise falls back to filename matching.
    Returns (content, filename).
    """
    if not os.path.exists(REF_DOCUMENTS_DIR):
        return "", ""
    
    files = glob.glob(os.path.join(REF_DOCUMENTS_DIR, "*.docx"))
    if not files:
        return "", ""
    
    
    if openai_key and os.path.exists(CHROMA_DB_DIR):
        try:
            embedding_function = OpenAIEmbeddings(api_key=openai_key)
            db = Chroma(persist_directory=CHROMA_DB_DIR, embedding_function=embedding_function)
            results = db.similarity_search(user_input, k=1)
            if results:
                best_doc = results[0]
                source = best_doc.metadata.get("source", "Unknown")
                return best_doc.page_content, os.path.basename(source)
        except Exception:
            pass  
    
    # Fallback: filename keyword matching
    best_file = None
    max_matches = 0
    normalized_input = user_input.lower()
    
    random.shuffle(files)
    
    for file_path in files:
        filename = os.path.basename(file_path)
        name_stem = os.path.splitext(filename)[0].lower()
        tokens = name_stem.replace('_', ' ').replace('-', ' ').split()
        matches = sum(1 for token in tokens if token in normalized_input and len(token) > 2)
        
        if matches > max_matches:
            max_matches = matches
            best_file = file_path
            
    selected_file = best_file if best_file else files[0]
    
    content = load_docx_content(selected_file)
    if content:
        return content, os.path.basename(selected_file)
    return "", ""


def ingest_exemplars_to_vector_db(openai_key: str) -> int:
    """
    Reads all .docx files in ref_documents, extracts text,
    and stores them in a local Vector DB.
    Returns number of documents ingested.
    """
    if not os.path.exists(REF_DOCUMENTS_DIR):
        return 0

    files = glob.glob(os.path.join(REF_DOCUMENTS_DIR, "*.docx"))
    if not files:
        return 0

    from langchain_core.documents import Document as LCDocument
    documents = []
    for file_path in files:
        content = load_docx_content(file_path)
        if content:
            documents.append(LCDocument(page_content=content, metadata={"source": file_path}))

    if not documents:
        return 0

    # Split text if cases are very long
    text_splitter = CharacterTextSplitter(chunk_size=2000, chunk_overlap=200)
    docs = text_splitter.split_documents(documents)

    # Create/Update Vector Store
    embedding_function = OpenAIEmbeddings(api_key=openai_key)
    
    # This saves the DB locally to disk
    Chroma.from_documents(
        documents=docs, 
        embedding=embedding_function, 
        persist_directory=CHROMA_DB_DIR
    )
    return len(docs)


MODEL_OPTIONS = {
    "OpenAI GPT-4o": {"provider": "openai", "id": "gpt-4o"},
    "Claude 3.5 Sonnet": {"provider": "anthropic", "id": "claude-sonnet-4-20250514"},
}


def get_model_config(label: str) -> dict:
    return MODEL_OPTIONS[label]


def convert_messages_for_anthropic(messages):
    system_parts = []
    convo_messages = []
    for msg in messages:
        role = msg.get("role")
        content = msg.get("content", "")
        if role == "system":
            if content:
                system_parts.append(content)
            continue
        if role in ("user", "assistant"):
            convo_messages.append({"role": role, "content": content})
    system_prompt = "\n\n".join(system_parts)
    return system_prompt, convo_messages


def stream_chat_chunks(model_config: dict, messages: list, temperature: float, max_tokens: int, anthropic_client: Optional[Anthropic]):
    provider = model_config["provider"]
    model_id = model_config["id"]
    clamped_temp = min(max(temperature, 0.0), 1.0)
    if provider == "openai":
        response_stream = openai.chat.completions.create(
            model=model_id,
            messages=messages,
            temperature=clamped_temp,
            max_tokens=max_tokens,
            stream=True
        )
        for chunk in response_stream:
            delta = chunk.choices[0].delta
            if not delta:
                continue
            delta_content = getattr(delta, "content", None)
            if delta_content is None and isinstance(delta, dict):
                delta_content = delta.get("content")
            if not delta_content:
                continue
            yield delta_content
        return
    if not anthropic_client:
        raise RuntimeError("Claude model selected but Anthropic API key is missing.")
    system_prompt, anthropic_messages = convert_messages_for_anthropic(messages)
    with anthropic_client.messages.stream(
        model=model_id,
        max_tokens=max_tokens,
        temperature=clamped_temp,
        system=system_prompt or "",
        messages=anthropic_messages,
    ) as stream:
        for event in stream:
            if event.type == "content_block_delta" and getattr(event.delta, "type", "") == "text_delta":
                yield event.delta.text
                
def strip_language_profile_heading(text: str) -> str:
    if not text:
        return ""
    stripped = text.lstrip()
    lines = stripped.splitlines()
    while lines:
        first_line = lines[0].strip()
        if LANGUAGE_PROFILE_HEADING_TEXT.lower() in first_line.lower():
            lines.pop(0)
            while lines and not lines[0].strip():
                lines.pop(0)
        else:
            break
    return "\n".join(lines).strip()

# Dedicated Language Profile agent helper
def generate_language_profile_section(
    case_text: str,
    user_prompt: str,
    model_config: dict,
    temperature: float,
    anthropic_client: Optional[Anthropic],
    on_chunk=None,
) -> str:
    profile_system_prompt = (
        "You are the Language Profile Agent for a speech-language pathology educator. "
        "Your only job is to translate a case narrative into a detailed language profile/communication observations summary."
    )
    exemplar = (
        "Language Sample / Communication Observations\n"
        "Collected through conversational tasks, picture description, narrative generation, and functional exchanges.\n"
        "Verbal Expression\n"
        "Fluent, effortless speech with intact prosody.\n"
        "Frequent semantic paraphasias, circumlocutions, occasional phonemic paraphasias, and occasional neologisms.\n"
        "Sentences are grammatically intact but often empty in content or tangential.\n"
        "Awareness of errors is inconsistent; spontaneous repairs are limited.\n"
        "Auditory Comprehension\n"
        "Difficulty understanding multi-step directions.\n"
        "Reduced accuracy for moderately complex yes/no questions.\n"
        "Difficulty understanding conversational discourse without visual supports or contextual cues.\n"
        "Repetition\n"
        "Moderately–severely impaired for multisyllabic words and sentence-level stimuli.\n"
        "Naming\n"
        "Moderate impairment in confrontation naming with semantic substitutions and occasional perseveration."
    )
    profile_user_prompt = f"""User request: {user_prompt}\n\nCase narrative: {case_text}\n\nProduce a Language Profile / Communication Observations section modeled on the exemplar below. Match the structure (intro plus subsections such as Verbal Expression, Auditory Comprehension, Repetition, Naming, etc.), but tailor every content line to the provided case details. Keep it concise, clinically rich, and coherent with the narrative.\n\nExemplar:\n{exemplar}"""
    messages = [
        {"role": "system", "content": profile_system_prompt},
        {"role": "user", "content": profile_user_prompt}
    ]
    profile_text = ""
    for delta_content in stream_chat_chunks(
        model_config=model_config,
        messages=messages,
        temperature=temperature,
        max_tokens=800,
        anthropic_client=anthropic_client,
    ):
        profile_text += delta_content
        if on_chunk:
            on_chunk(profile_text)
    return profile_text.strip()

# Dedicated RTSS agent helper
def generate_rtss_section(
    case_text: str,
    user_prompt: str,
    model_config: dict,
    temperature: float,
    anthropic_client: Optional[Anthropic],
    on_chunk=None,
) -> str:
    rtss_system_prompt = (
        "You are the RTSS Agent, specializing in Rehabilitation Treatment Specification System "
        "frameworks for speech-language pathology. Your sole task is to translate a given "
        "SLP case narrative into an RTSS-aligned plan that stays perfectly coherent with the "
        "provided case details."
    )
    rtss_user_prompt = f"""User request: {user_prompt}\n\nCase narrative: {case_text}\n\nCreate an RTSS section that includes: \n- Target (participation, impairment, or contextual focus) tied to the case goals.\n- Ingredients (specific clinician actions/techniques) with enough detail for another SLP to reproduce.\n- Mechanisms of action explaining why those ingredients are expected to work.\n- Dosage/parameters (session length, frequency, cues, materials).\n- Expected short-term markers of progress.\nFormat as concise markdown with clear subheadings. Do not contradict the narrative; infer only what is supported by it."""
    messages = [
        {"role": "system", "content": rtss_system_prompt},
        {"role": "user", "content": rtss_user_prompt}
    ]
    section_text = ""
    for delta_content in stream_chat_chunks(
        model_config=model_config,
        messages=messages,
        temperature=temperature,
        max_tokens=800,
        anthropic_client=anthropic_client,
    ):
        section_text += delta_content
        if on_chunk:
            on_chunk(section_text)
    return section_text.strip()

# Page configuration
st.set_page_config(
    page_title="Clinical Case Study Generator",
    page_icon="📚",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS
st.markdown("""
<style>
    .main-header {
        font-size: 2.5rem;
        font-weight: bold;
        color: #4F46E5;
        text-align: center;
    }
    .sub-header {
        font-size: 1rem;
        color: #6B7280;
        margin-bottom: 2rem;
    }
    .stTextInput > div > div > input {
        border-radius: 10px;
    }
    .chat-message {
        padding: 1.5rem;
        border-radius: 10px;
        margin-bottom: 1rem;
        color: #1F2937;
        position: relative;
    }
    .user-message {
        background-color: #EEF2FF;
        border-left: 4px solid #4F46E5;
    }
    .assistant-message {
        background-color: #F9FAFB;
        border-left: 4px solid #10B981;
    }
    /* Force dark text in chat messages regardless of theme */
    .chat-message * {
        color: #1F2937 !important;
    }
    .chat-message strong {
        color: #111827 !important;
    }
</style>
""", unsafe_allow_html=True)


def render_user_message(content: str) -> str:
    return f'<div class="chat-message user-message"><strong>👤 You:</strong><br>{content}</div>'

def render_assistant_message(content: str) -> str:
    """Render assistant message HTML."""
    return f'<div class="chat-message assistant-message"><strong>🤖 Assistant:</strong><br>{content}</div>'

def render_copy_button(content: str):
    """Render a functional copy button using st.components.html()."""
    # Use json.dumps for JavaScript string literal, then html.escape for HTML attribute context
    js_string = json.dumps(content)
    escaped_js_string = html.escape(js_string)
    html_code = f'''
    <button id="copyBtn" style="
        background-color: #E5E7EB;
        border: none;
        border-radius: 5px;
        padding: 5px 10px;
        cursor: pointer;
        font-size: 0.8rem;
        color: #374151;
        transition: background-color 0.2s;
    " onmouseover="this.style.backgroundColor='#D1D5DB'" 
       onmouseout="this.style.backgroundColor='#E5E7EB'"
       onclick="
        var text = {escaped_js_string};
        var btn = document.getElementById('copyBtn');
        function showCopied() {{
            btn.textContent = '✓ Copied!';
            setTimeout(function() {{ btn.textContent = '📋 Copy'; }}, 2000);
        }}
        navigator.clipboard.writeText(text).then(showCopied).catch(function(err) {{
            // Fallback for older browsers
            var textarea = document.createElement('textarea');
            textarea.value = text;
            textarea.style.position = 'fixed';
            textarea.style.opacity = '0';
            document.body.appendChild(textarea);
            textarea.select();
            try {{
                document.execCommand('copy');
                showCopied();
            }} catch(e) {{
                console.error('Copy failed:', e);
            }}
            document.body.removeChild(textarea);
        }});
    ">📋 Copy</button>
    '''
    components.html(html_code, height=35)

# Get API keys from secrets (keys validated later based on provider selection)
openai_api_key = st.secrets["OPENAI_API_KEY"] if "OPENAI_API_KEY" in st.secrets else None
if openai_api_key:
    openai.api_key = openai_api_key
anthropic_api_key = st.secrets["ANTHROPIC_API_KEY"] if "ANTHROPIC_API_KEY" in st.secrets else None
anthropic_client: Optional[Anthropic] = Anthropic(api_key=anthropic_api_key) if anthropic_api_key else None

# Initialize session state
if 'messages' not in st.session_state:
    st.session_state.messages = []
if 'rtss_enabled' not in st.session_state:
    st.session_state.rtss_enabled = False
if 'language_profile_enabled' not in st.session_state:
    st.session_state.language_profile_enabled = False

# Sidebar
with st.sidebar:
    st.markdown("### 📖 Quick Prompts")
    
    quick_prompts = [
        "create a moderate Broca's aphasia case for learning initial assessment",
        "create a dementia case for treatment planning",
        "Demonstration of collaborative goal setting with conversation scripts",
        "Demonstration of motivational interviewing with conversation scripts",
        "Demonstration of a specific treatment technique with conversation scripts"
    ]
    dropdown_options = ["Select a quick prompt…"] + quick_prompts
    selected_prompt = st.selectbox("Choose a quick prompt", dropdown_options, index=0)
    if st.button("Send Quick Prompt", use_container_width=True, disabled=selected_prompt == dropdown_options[0]):
        st.session_state["pending_user_input"] = selected_prompt
        st.rerun()
    
    st.markdown("---")
    st.markdown("### ➕ Additional Content")
    st.session_state.language_profile_enabled = st.checkbox(
        "Generate Language Profile section",
        value=st.session_state.get("language_profile_enabled", False),
        help="Check to append a detailed Language Profile / Communication Observations summary modeled on clinical documentation."
    )
    st.session_state.rtss_enabled = st.checkbox(
        "Generate RTSS section",
        value=st.session_state.get("rtss_enabled", False),
        help="Check to append a Rehabilitation Treatment Specification System plan after each case."
    )
    
    st.markdown("---")
    st.markdown("### ⚙️ Settings")
    difficulty = "Intermediate (Graduate)"
    include_assessment = True
    
    # Model selection
    model_labels = list(MODEL_OPTIONS.keys())
    selected_model_label = st.selectbox(
        "Model",
        model_labels,
        index=0,
        help="Choose between the default OpenAI model and the integrated Claude model"
    )
    
    # Temperature
    temperature = st.slider(
        "Creativity",
        min_value=0.0,
        max_value=1.0,
        value=0.7,
        step=0.1,
        help="Higher values make output more creative, lower values more focused"
    )
    # Clear conversation button
    if st.button("🗑️ Clear Conversation", use_container_width=True):
        st.session_state.messages = []
        st.rerun()
    
    st.markdown("---")

model_config = get_model_config(selected_model_label)
if model_config["provider"] == "openai" and not openai_api_key:
    st.error("⚠️ OpenAI model selected but OPENAI_API_KEY is not configured in .streamlit/secrets.toml")
    st.stop()
if model_config["provider"] == "anthropic" and not anthropic_client:
    st.error("⚠️ Claude model selected but ANTHROPIC_API_KEY is not configured in .streamlit/secrets.toml")
    st.stop()

# Main content
st.markdown('<div class="main-header">📚 Clinical Case Study Generator</div>', unsafe_allow_html=True)
st.markdown(
    """
    <div style='text-align: center; color: #6B7280; padding: 0;'>
        <p>An AI-powered educational case study generator for neurogenic communication disorders</p>
    </div>
    """,
    unsafe_allow_html=True
)

live_user_placeholder = None
live_assistant_placeholder = None
conversation_container = st.container()
with conversation_container:
    for message in st.session_state.messages:
        if message["role"] == "user":
            st.markdown(render_user_message(message["content"]), unsafe_allow_html=True)
        else:
            st.markdown(render_assistant_message(message["content"]), unsafe_allow_html=True)
            render_copy_button(message["content"])
    live_user_placeholder = st.empty()
    live_assistant_placeholder = st.empty()

# Chat input (process queued quick prompts before showing input box)
queued_input = st.session_state.pop("pending_user_input", None)
user_input = st.chat_input("Describe the case study you'd like to create...")
pending_input = queued_input or user_input

if pending_input and live_user_placeholder and live_assistant_placeholder:
    st.session_state.messages.append({"role": "user", "content": pending_input})
    live_user_placeholder.markdown(render_user_message(pending_input), unsafe_allow_html=True)
    normalized_input = pending_input.lower()
    is_initial_assessment = "initial assessment" in normalized_input
    transfer_keywords = ["transfer", "new setting", "transition", "new facility", "new clinic","discharge to"]
    is_transfer_initial = is_initial_assessment and any(keyword in normalized_input for keyword in transfer_keywords)
    rtss_keywords = ["rtss", "rehabilitation treatment specification", "treatment specification"]
    rtss_requested = any(keyword in normalized_input for keyword in rtss_keywords)
    language_profile_keywords = ["language profile", "communication observations", "verbal expression profile"]
    language_profile_requested = any(keyword in normalized_input for keyword in language_profile_keywords)
    
    # Load a reference case if available, prioritizing semantic matches
    reference_content, reference_name = get_best_matching_exemplar(pending_input, openai_api_key)
    reference_instruction = ""
    if reference_content:
        reference_instruction = f"""
### Reference Style Guide
Below is an example of an ideal case study. Use it ONLY for structural and stylistic inspiration (tone, depth of detail, phrasing). 
DO NOT copy the patient details, diagnosis, or specific scenario from this reference. 
Apply this high-quality clinical writing style to the specific case requested by the user.

<reference_exemplar>
{reference_content}
</reference_exemplar>
"""
    
    # Create system prompt based on settings
    difficulty_map = {
        "Beginner (Undergraduate)": "beginner",
        "Intermediate (Graduate)": "intermediate",
        "Advanced (Clinical Fellows)": "advanced"
    }
    guiding_focus = "Focus these questions on determining initial assessment priorities and selecting appropriate standardized and informal tools." if is_initial_assessment else "Use these questions to spark discussion about differential diagnosis, intervention planning, and clinical reasoning."
    if is_initial_assessment and not is_transfer_initial:
        language_sample_note = "Ensure the language sample reflects spontaneous speech observable prior to formal testing."
    else:
        language_sample_note = ""
    initial_assessment_guidance = ""
    if is_initial_assessment:
        if is_transfer_initial:
            initial_assessment_guidance = (
                "\nWhen the user signals an initial assessment during a transfer of care, summarize what prior SLPs accomplished, the patient's recovery trajectory, and documentation the new setting would inherit. Highlight unanswered assessment questions for the new site and clarify what still requires standardized and informal testing."
            )
        else:
            initial_assessment_guidance = (
                "\nWhen the user signals an initial assessment for a new onset with no prior SLP involvement, limit the case to referral information, medical history, observable behaviors, and collateral reports available before assessments are administered. Do not invent completed SLP assessment results yet; instead, describe what needs to be investigated and why."
            )
    
    system_prompt = f"""You are an expert clinical educator specializing in speech-language pathology. Your role is to help university instructors create high-quality, realistic case studies for their students.
{reference_instruction}
Always craft a single cohesive narrative case study (paragraph style, not bullet points) that weaves together (and do not pre-answer or give instructions for the guiding questions here):
- Patient demographics and background
- Detailed medical history and etiology
- Presenting symptoms and characteristics
- Assessment results (formal and informal) with interpretation
- Clinical observations and differential diagnosis considerations
- Treatment recommendations, prognosis.

Keep the narrative observational and exploratory; reserve any prioritization, tool selection, or explicit answers for the Guiding Questions section only.

After the narrative, append two explicit sections. Bold the headers exactly as shown ("**Language Sample**" and "**Guiding Questions**") so they stand out visually:
1. **Language Sample** — provide a short quoted transcript (4–6 sentences) that captures the client's spontaneous speech. Match the cadence of this reference format: "Well the, uh, the little… the little cookers are spinning up there and she's trying to wash the plates but the water's all, all floofing out. And the boy, he's, he's grabbing the stool cause he wants the cookie. They're having a good time, I think, and the mother doesn't know the window is, is, uh, smiling. It's pretty noisy in that room." Use it only as stylistic guidance; compose a fresh sample aligned with the case's disorder-specific features every time so it reflects the symptoms described in the user input. {language_sample_note}
2. **Guiding Questions** — provide discussion questions instructors can use in class. {guiding_focus}

Keep the entire response focused on the case (no general tips). Adjust complexity to {difficulty_map[difficulty]} level, ensure clinical accuracy, and use professional terminology appropriate for graduate-level speech-language pathology education.{initial_assessment_guidance}"""
    
    # Prepare messages for API
    api_messages = [{"role": "system", "content": system_prompt}] + [
        {"role": m["role"], "content": m["content"]} 
        for m in st.session_state.messages
    ]
    
    # Stream model response so content appears as it arrives
    try:
        assistant_message = ""
        with st.spinner("Generating case study..."):
            for delta_content in stream_chat_chunks(
                model_config=model_config,
                messages=api_messages,
                temperature=temperature,
                max_tokens=2000,
                anthropic_client=anthropic_client,
            ):
                assistant_message += delta_content
                live_assistant_placeholder.markdown(
                    render_assistant_message(assistant_message),
                    unsafe_allow_html=True
                )
        
        base_message = assistant_message

        # Generate Language Profile section if requested
        if st.session_state.get("language_profile_enabled", False) or language_profile_requested:
            try:
                with st.spinner("Drafting language profile..."):
                    def update_language_profile(section_partial: str) -> None:
                        sanitized_partial = strip_language_profile_heading(section_partial)
                        live_assistant_placeholder.markdown(
                            render_assistant_message(
                                f"{base_message}\n\n{LANGUAGE_PROFILE_HEADING}\n{sanitized_partial}"
                            ),
                            unsafe_allow_html=True
                        )
                    language_profile_section = generate_language_profile_section(
                        case_text=base_message,
                        user_prompt=pending_input,
                        model_config=model_config,
                        temperature=temperature,
                        anthropic_client=anthropic_client,
                        on_chunk=update_language_profile
                    )
                language_profile_section = strip_language_profile_heading(language_profile_section)
                if language_profile_section:
                    assistant_message = (
                        f"{base_message}\n\n{LANGUAGE_PROFILE_HEADING}\n{language_profile_section}"
                    )
                    live_assistant_placeholder.markdown(
                        render_assistant_message(assistant_message),
                        unsafe_allow_html=True
                    )
                    base_message = assistant_message
            except Exception as lp_error:
                st.warning(f"Language profile agent skipped due to error: {lp_error}")

        # Generate RTSS section with dedicated agent and append to response
        if st.session_state.get("rtss_enabled", False) or rtss_requested:
            try:
                with st.spinner("Synthesizing RTSS plan..."):
                    base_message = assistant_message
                    def update_rtss(section_partial: str) -> None:
                        live_assistant_placeholder.markdown(
                            render_assistant_message(
                                f"{base_message}\n\n{section_partial}"
                            ),
                            unsafe_allow_html=True
                        )
                    rtss_section = generate_rtss_section(
                        case_text=base_message,
                        user_prompt=pending_input,
                        model_config=model_config,
                        temperature=temperature,
                        anthropic_client=anthropic_client,
                        on_chunk=update_rtss
                    )
                if rtss_section:
                    assistant_message = (
                        f"{base_message}\n\n{rtss_section}"
                    )
                    live_assistant_placeholder.markdown(
                        render_assistant_message(assistant_message),
                        unsafe_allow_html=True
                    )
            except Exception as rtss_error:
                st.warning(f"RTSS agent skipped due to error: {rtss_error}")
        
        st.session_state.messages.append({"role": "assistant", "content": assistant_message})
        st.rerun()
        
    except Exception as e:
        st.error(f"❌ Error: {str(e)}")
        st.info("Please check your API configuration and try again.")