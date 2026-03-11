"""
All RAG-related logic: document loading, vector-DB ingestion,
and exemplar retrieval (semantic + keyword fallback).
"""

import os
import glob
import random
from typing import Optional

from docx import Document as DocxDocument
from config import REF_DOCUMENTS_DIR, CHROMA_DB_DIR

# Document loading
def load_docx_content(file_path: str) -> str:
    """Extract plain text from a .docx file."""
    try:
        doc = DocxDocument(file_path)
        return "\n".join(para.text for para in doc.paragraphs if para.text.strip())
    except Exception:
        return ""


# Vector DB ingestion
def ingest_exemplars_to_vector_db(openai_key: str) -> int:
    """
    Read all .docx files in ref_documents, extract text, and store them in
    a local ChromaDB vector store.  Returns the number of chunks ingested.
    """
    if not os.path.exists(REF_DOCUMENTS_DIR):
        return 0

    files = glob.glob(os.path.join(REF_DOCUMENTS_DIR, "*.docx"))
    if not files:
        return 0

    from langchain_core.documents import Document as LCDocument
    from langchain_text_splitters import CharacterTextSplitter
    from langchain_openai import OpenAIEmbeddings
    from langchain_community.vectorstores import Chroma

    documents = []
    for file_path in files:
        content = load_docx_content(file_path)
        if content:
            documents.append(
                LCDocument(page_content=content, metadata={"source": file_path})
            )

    if not documents:
        return 0

    text_splitter = CharacterTextSplitter(chunk_size=2000, chunk_overlap=200)
    docs = text_splitter.split_documents(documents)

    embedding_function = OpenAIEmbeddings(api_key=openai_key)
    Chroma.from_documents(
        documents=docs,
        embedding=embedding_function,
        persist_directory=CHROMA_DB_DIR,
    )
    return len(docs)


# Exemplar retrieval
def get_best_matching_exemplar(
    user_input: str,
    openai_key: Optional[str],
) -> tuple[str, str]:
    """
    Return the (content, filename) of the best-matching reference case.

    Priority:
      1. Semantic search via ChromaDB (when openai_key is available and DB exists).
      2. Keyword match against filenames.
      3. Random file from the directory.
    """
    if not os.path.exists(REF_DOCUMENTS_DIR):
        return "", ""

    files = glob.glob(os.path.join(REF_DOCUMENTS_DIR, "*.docx"))
    if not files:
        return "", ""

    # --- Semantic search ---
    if openai_key and os.path.exists(CHROMA_DB_DIR):
        try:
            from langchain_openai import OpenAIEmbeddings
            from langchain_community.vectorstores import Chroma

            embedding_function = OpenAIEmbeddings(api_key=openai_key)
            db = Chroma(
                persist_directory=CHROMA_DB_DIR,
                embedding_function=embedding_function,
            )
            results = db.similarity_search(user_input, k=1)
            if results:
                best_doc = results[0]
                source = best_doc.metadata.get("source", "Unknown")
                return best_doc.page_content, os.path.basename(source)
        except Exception:
            pass  # Fall through to keyword matching

    # --- Keyword / filename matching ---
    best_file = None
    max_matches = 0
    normalized_input = user_input.lower()
    random.shuffle(files)

    for file_path in files:
        name_stem = os.path.splitext(os.path.basename(file_path))[0].lower()
        tokens = name_stem.replace("_", " ").replace("-", " ").split()
        matches = sum(
            1 for token in tokens if token in normalized_input and len(token) > 2
        )
        if matches > max_matches:
            max_matches = matches
            best_file = file_path

    selected_file = best_file if best_file else files[0]
    content = load_docx_content(selected_file)
    if content:
        return content, os.path.basename(selected_file)
    return "", ""